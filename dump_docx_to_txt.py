from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    import argparse

    ap = argparse.ArgumentParser()
    ap.add_argument("docx_path", type=str)
    ap.add_argument("out_txt", type=str)
    args = ap.parse_args()

    doc = Document(args.docx_path)
    out = Path(args.out_txt)

    lines: list[str] = []
    lines.append(f"FILE: {args.docx_path}")
    lines.append(f"PARAGRAPHS: {len(doc.paragraphs)}")
    lines.append("")

    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "None"
        txt = (p.text or "").replace("\t", " ").strip()
        if not txt:
            continue
        lines.append(f"{i:04d} [{style}] {txt}")

    out.write_text("\n".join(lines), encoding="utf-8")
    print("Wrote", out)


if __name__ == "__main__":
    main()

