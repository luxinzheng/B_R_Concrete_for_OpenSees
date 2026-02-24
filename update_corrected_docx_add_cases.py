from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Cm


SRC = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper_corrected.docx")
OUT = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper_corrected_v2.docx")
BACKUP = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper_corrected.bak.docx")

FIG_UNIAX = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_uniaxial.png")
FIG_WALL = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_wall.png")
FIG_BEAM = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_beam.png")
FIG_PLATE = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_plate.png")


def remove_paragraph(paragraph: Paragraph) -> None:
    p: CT_P = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None  # type: ignore


def insert_paragraph_before(anchor: Paragraph, text: str, style: Optional[str] = None) -> Paragraph:
    p = anchor.insert_paragraph_before(text)
    if style:
        p.style = style
    return p


def insert_picture(doc: Document, anchor: Paragraph, image_path: Path, width_cm: float = 15.5) -> None:
    # Insert a blank paragraph before anchor, then add picture there.
    pic_p = anchor.insert_paragraph_before("")
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_first_paragraph_index(doc: Document, startswith: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    raise ValueError(f"Cannot find paragraph starting with: {startswith!r}")


def replace_abstract(doc: Document) -> None:
    # Append one sentence to abstract to reflect expanded verification cases.
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("摘要："):
            if "悬臂梁" in t and "RC板" in t:
                return
            add = "此外，验证工作扩展为单轴、剪力墙、悬臂梁与RC板等4类结构级算例，并进一步在多片剪力墙与多层壳单元算例上进行软件级稳定性校核，增强了结论的工程代表性与可复现性。"
            p.text = t + add
            return
    # If not found, silently skip.


def rebuild_section_4(doc: Document) -> None:
    # Identify section 4 range: from '4 验证设计及结果' (Subtitle) to just before '5 讨论' (Subtitle)
    i4 = find_first_paragraph_index(doc, "4 ")
    i5 = find_first_paragraph_index(doc, "5 ")

    # Remove paragraphs in [i4, i5)
    # Need to remove from end to start to preserve indices
    for p in list(doc.paragraphs[i4:i5])[::-1]:
        remove_paragraph(p)

    # After removal, paragraphs list is refreshed; find anchor where to insert (the current '5 讨论')
    anchor = doc.paragraphs[find_first_paragraph_index(doc, "5 ")]

    # Insert new Section 4 content before anchor
    insert_paragraph_before(anchor, "4 验证设计及结果", style="Subtitle")

    insert_paragraph_before(anchor, "4.1 验证体系与评价指标", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "验证采用“单元级—结构级—软件级”三级体系。单元级通过 standalone Fortran 程序直接调用 PSUMAT 接口，"
        "在不依赖 OpenSees 可执行文件的前提下开展回归测试；结构级将 forumat.f90 编译进 OpenSees 后，"
        "采用与参考模型一致的算例与加载协议对比响应；软件级进一步选取多片剪力墙与多层壳单元等算例，"
        "检验大变形循环加载下的收敛性与数值稳定性。剪力墙算例中，基底剪力取底部约束节点反力之和 "
        "V=ΣRi(i=1…5)，控制位移取加载点水平位移 u，比较 V-u 滞回曲线峰值、对称性与收敛步数等指标。",
        style="Normal",
    )

    insert_paragraph_before(anchor, "4.2 单元级材料测试（5个）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "依据开发记录，单元级材料测试覆盖5条典型路径：单轴压缩（至 ε=-0.003）、单轴拉伸（至 ε=+0.0005）、"
        "纯剪切（γ=0.005）、循环压拉（压缩→卸载→拉伸→再压缩）与四阶段循环（压→卸→拉→再压）。"
        "上述测试均顺利完成且无失败步，Saenz 型压缩曲线峰后软化、拉伸开裂软化与循环路径均满足预期，"
        "为后续结构级算例提供了可追溯的单元级正确性保证。",
        style="Normal",
    )

    insert_paragraph_before(anchor, "4.3 验证套件（4类，与参考模型对比）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "为验证集成本构在结构响应层面的有效性，建立了4类对比验证套件：单轴（压缩/拉伸/循环）、"
        "悬臂剪力墙循环加载、RC悬臂梁循环加载与RC板弯曲。对比结果表明：B&R 与参考模型在峰值、软化段"
        "与滞回形态上总体吻合，其中剪力墙峰值底部剪力略低约 5%～10%，梁算例弹性段几乎完全重合，"
        "开裂后 B&R 略软；RC板算例两者基本重合。对应对比图分别见图3～图6。",
        style="Normal",
    )

    # Captions + pictures
    p_fig3 = insert_paragraph_before(
        anchor,
        "图 3 单轴压缩与拉伸应力-应变曲线（集成模型与参考模型）\n"
        "Figure 3 Uniaxial compressive and tensile stress-strain curves (integrated model vs reference)",
        style="Normal",
    )
    if FIG_UNIAX.exists():
        insert_picture(doc, anchor=p_fig3, image_path=FIG_UNIAX)

    p_fig4 = insert_paragraph_before(
        anchor,
        "图 4 剪力墙侧向力-位移滞回曲线（集成模型与参考模型）\n"
        "Figure 4 Shear wall lateral force-displacement hysteresis (integrated model vs reference)",
        style="Normal",
    )
    if FIG_WALL.exists():
        insert_picture(doc, anchor=p_fig4, image_path=FIG_WALL)

    p_fig5 = insert_paragraph_before(
        anchor,
        "图 5 悬臂梁循环加载力-位移对比（集成模型与参考模型）\n"
        "Figure 5 Cantilever beam cyclic response comparison (integrated model vs reference)",
        style="Normal",
    )
    if FIG_BEAM.exists():
        insert_picture(doc, anchor=p_fig5, image_path=FIG_BEAM)

    p_fig6 = insert_paragraph_before(
        anchor,
        "图 6 RC板弯曲荷载-挠度对比（集成模型与参考模型）\n"
        "Figure 6 RC slab bending load-deflection comparison (integrated model vs reference)",
        style="Normal",
    )
    if FIG_PLATE.exists():
        insert_picture(doc, anchor=p_fig6, image_path=FIG_PLATE)

    insert_paragraph_before(anchor, "4.4 多片剪力墙与多层壳单元软件级校核", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "除上述验证套件外，进一步选取多片剪力墙与多层壳单元（Multi-layer Shell）算例开展软件级稳定性校核。"
        "开发记录显示：sw1-1 剪力墙可完成 500/500 步、0 失败；sw2-1 剪力墙在大位移后期完成约 90% 进度，"
        "出现矩阵奇异相关的收敛困难；Multi-layer Shell 算例可完成全部 947/947 步、0 失败。该类算例更贴近"
        "工程分析中“大变形+循环+耦合单元”的使用情景，可用于检验材料切线处理与求解器耦合的鲁棒性。",
        style="Normal",
    )

    insert_paragraph_before(anchor, "4.5 验证小结（表 3）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "表 3 汇总单元级材料测试、4类验证套件与软件级校核的主要指标与结论。总体而言，新增校验算例覆盖了"
        "“材料层—构件层—结构层—软件层”的多尺度验证链条，显著增强了结论的工程代表性。",
        style="Normal",
    )

    # Table 3
    tcap = insert_paragraph_before(
        anchor,
        "表 3 验证结果汇总\nTable 3 Summary of verification results",
        style="Normal",
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "类别"
    hdr[1].text = "算例/路径"
    hdr[2].text = "关键指标"
    hdr[3].text = "结论"

    def add_row(cat, case, metric, concl):
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = case
        row[2].text = metric
        row[3].text = concl

    add_row("单元级材料测试", "单轴压缩/拉伸/纯剪切/循环压拉/四阶段循环", "0 失败；软化/滞回路径符合预期", "通过")
    add_row("结构级验证套件", "单轴（压缩/拉伸/循环）", "峰值与软化段与参考吻合", "通过")
    add_row("结构级验证套件", "剪力墙循环加载", "滞回形态一致；峰值略低约 5%～10%", "通过")
    add_row("结构级验证套件", "悬臂梁循环加载", "弹性段几乎重合；开裂后略软", "通过")
    add_row("结构级验证套件", "RC板弯曲", "与参考基本重合", "通过")
    add_row("软件级校核", "sw1-1 多片剪力墙", "500/500 步，0 失败", "通过")
    add_row("软件级校核", "sw2-1 多片剪力墙", "≈452/500 步，后期奇异/收敛困难", "部分通过")
    add_row("软件级校核", "Multi-layer Shell", "947/947 步，0 失败", "通过")

    # Move table to be right after caption paragraph by inserting it before anchor (Word appends at end).
    # Workaround: insert a marker paragraph and move table XML before it.
    marker = anchor.insert_paragraph_before("")
    marker._element.addprevious(table._element)
    # Keep caption just before table
    # (caption is already inserted; ensure order: caption -> table -> marker removed)
    remove_paragraph(marker)

    # spacer before section 5
    insert_paragraph_before(anchor, "")


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    # backup original
    if not BACKUP.exists():
        BACKUP.write_bytes(SRC.read_bytes())

    doc = Document(str(SRC))
    replace_abstract(doc)
    rebuild_section_4(doc)
    doc.save(str(OUT))
    print("Wrote", OUT)
    print("Backup", BACKUP)


if __name__ == "__main__":
    main()

