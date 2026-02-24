from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    import argparse

    ap = argparse.ArgumentParser()
    ap.add_argument("docx_path", type=str)
    ap.add_argument("--max", type=int, default=400)
    args = ap.parse_args()

    p = Path(args.docx_path)
    doc = Document(str(p))
    print("FILE:", p)
    print("PARAGRAPHS:", len(doc.paragraphs))
    print("-" * 80)
    shown = 0
    for i, para in enumerate(doc.paragraphs):
        txt = (para.text or "").strip().replace("\t", " ")
        if not txt:
            continue
        style = para.style.name if para.style else "None"
        is_heading = style.lower().startswith("heading") or style.lower() in {"title", "subtitle"}
        is_key = any(k in txt for k in ("摘要", "关键词", "Abstract", "Keywords", "引言", "结论", "参考文献", "致谢"))
        if is_heading or is_key or txt[:2].isdigit():
            print(f"{i:04d} [{style}] {txt[:160]}")
            shown += 1
            if shown >= args.max:
                break


if __name__ == "__main__":
    main()

