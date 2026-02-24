from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.shared import Cm


SRC = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper_corrected_v2.docx")
OUT = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper_corrected_v3.docx")

FIG_UNIAX = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_uniaxial.png")
FIG_WALL = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_wall.png")

FIG_SW1 = Path(r"e:\Basic\Concrete_Model\ADINA\compare_sw1-1.png")
FIG_SW2 = Path(r"e:\Basic\Concrete_Model\ADINA\compare_sw2-1.png")
FIG_ML = Path(r"e:\Basic\Concrete_Model\ADINA\compare_Multi-layer_Shell.png")
FIG_ALL = Path(r"e:\Basic\Concrete_Model\ADINA\all_cases_comparison_0818.png")


def remove_paragraph(paragraph: Paragraph) -> None:
    p: CT_P = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None  # type: ignore


def insert_paragraph_before(anchor: Paragraph, text: str, style: Optional[str] = None) -> Paragraph:
    p = anchor.insert_paragraph_before(text)
    if style:
        p.style = style
    return p


def insert_picture(anchor: Paragraph, image_path: Path, width_cm: float = 15.5) -> None:
    pic_p = anchor.insert_paragraph_before("")
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_first_paragraph_index(doc: Document, startswith: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    raise ValueError(f"Cannot find paragraph starting with: {startswith!r}")


def normalize_abstract(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("摘要："):
            # remove any old sentence mentioning beam/plate
            t2 = re.sub(r"此外，验证工作扩展为.*?可复现性。", "", t)
            if t2 == t:
                # try remove any clause with 梁/板
                t2 = t
                t2 = t2.replace("悬臂梁与RC板等4类结构级算例，并", "")
                t2 = t2.replace("悬臂梁与RC板等4类结构级算例", "")
            add = "此外，验证工作扩展为单轴与剪力墙的结构级对比校验，并在多片剪力墙（sw1-1、sw2-1）与多层壳单元（Multi-layer Shell）算例上开展软件级稳定性校核，增强了结论的工程代表性与可复现性。"
            if add not in t2:
                t2 = t2 + add
            p.text = t2
            return


def rebuild_section_4(doc: Document) -> None:
    i4 = find_first_paragraph_index(doc, "4 ")
    i5 = find_first_paragraph_index(doc, "5 ")

    # remove old section 4 content
    for p in list(doc.paragraphs[i4:i5])[::-1]:
        remove_paragraph(p)

    anchor = doc.paragraphs[find_first_paragraph_index(doc, "5 ")]

    insert_paragraph_before(anchor, "4 验证设计及结果", style="Subtitle")

    insert_paragraph_before(anchor, "4.1 验证体系与评价指标", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "验证采用“单元级—结构级—软件级”三级体系。单元级通过 standalone Fortran 程序直接调用 PSUMAT 接口，"
        "在不依赖 OpenSees 可执行文件的前提下开展回归测试；结构级将 forumat.f90 编译进 OpenSees 后，"
        "采用与参考模型一致的算例与加载协议对比响应；软件级进一步选取多片剪力墙与多层壳单元等算例，"
        "检验大变形循环加载下的收敛性与数值稳定性。剪力墙算例中，基底剪力取底部约束节点反力之和 "
        "V=ΣRi(i=1…5)，控制位移取加载点水平位移 u，比较 V-u 滞回曲线峰值、对称性与收敛步数等指标。",
        style="Normal",
    )

    insert_paragraph_before(anchor, "4.2 单元级材料测试（5个）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "依据最新开发记录，单元级材料测试覆盖5条典型路径：单轴压缩（至 ε=-0.003）、单轴拉伸（至 ε=+0.0005）、"
        "纯剪切（γ=0.005）、循环压拉（压缩→卸载→拉伸→再压缩）与四阶段循环（压→卸→拉→再压）。"
        "上述测试均顺利完成且无失败步，为后续结构级与软件级算例提供了可追溯的单元级正确性保证。",
        style="Normal",
    )

    insert_paragraph_before(anchor, "4.3 结构级对比校验（单轴+剪力墙）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "结构级校验聚焦两类最具代表性的平面应力应用场景：单轴（压缩/拉伸/循环）与悬臂剪力墙循环加载。"
        "对比结果表明：B&R 与参考模型在单轴峰值与软化段吻合良好；剪力墙滞回环形态一致，B&R 峰值底部剪力"
        "略低约 5%～10%，包络线与退化趋势吻合。对应对比图见图3～图4。",
        style="Normal",
    )

    p_fig3 = insert_paragraph_before(
        anchor,
        "图 3 单轴压缩与拉伸应力-应变曲线（集成模型与参考模型）\n"
        "Figure 3 Uniaxial compressive and tensile stress-strain curves (integrated model vs reference)",
        style="Normal",
    )
    if FIG_UNIAX.exists():
        insert_picture(p_fig3, FIG_UNIAX)

    p_fig4 = insert_paragraph_before(
        anchor,
        "图 4 剪力墙侧向力-位移滞回曲线（集成模型与参考模型）\n"
        "Figure 4 Shear wall lateral force-displacement hysteresis (integrated model vs reference)",
        style="Normal",
    )
    if FIG_WALL.exists():
        insert_picture(p_fig4, FIG_WALL)

    insert_paragraph_before(anchor, "4.4 多片剪力墙与多层壳单元软件级校核（含对比图与成图脚本）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "为检验本构在更贴近工程使用情景（大变形+循环+耦合单元/求解器）的稳定性，进一步选取多片剪力墙"
        "（sw1-1、sw2-1）与多层壳单元（Multi-layer Shell）算例开展软件级校核。最新记录显示：sw1-1 可完成"
        "500/500 步、0 失败；sw2-1 在大位移后期完成约 90% 进度后出现矩阵奇异相关的收敛困难；Multi-layer Shell"
        "可完成 947/947 步、0 失败。三算例与参考模型的对比图分别见图5～图7。"
        "对应成图与对比分析脚本已随工程目录提供，主要包括："
        "plot_three_cases.py（多算例汇总对比）、plot_0901_comparison.py（sw1-1 对比成图）、"
        "plot_ml_comparison.py（Multi-layer Shell 对比成图）、plot_wall_comparison.py（墙体滞回对比），"
        "以及 verify/analyze_wall*.py（校验套件数据处理）。",
        style="Normal",
    )

    # Figures 5-7
    p_fig5 = insert_paragraph_before(
        anchor,
        "图 5 sw1-1 剪力墙对比（集成模型与参考模型）\n"
        "Figure 5 sw1-1 shear wall comparison (integrated model vs reference)",
        style="Normal",
    )
    if FIG_SW1.exists():
        insert_picture(p_fig5, FIG_SW1)

    p_fig6 = insert_paragraph_before(
        anchor,
        "图 6 sw2-1 剪力墙对比（集成模型与参考模型）\n"
        "Figure 6 sw2-1 shear wall comparison (integrated model vs reference)",
        style="Normal",
    )
    if FIG_SW2.exists():
        insert_picture(p_fig6, FIG_SW2)

    p_fig7 = insert_paragraph_before(
        anchor,
        "图 7 Multi-layer Shell 对比（集成模型与参考模型）\n"
        "Figure 7 Multi-layer Shell comparison (integrated model vs reference)",
        style="Normal",
    )
    if FIG_ML.exists():
        insert_picture(p_fig7, FIG_ML)

    # Optional summary figure (all cases)
    if FIG_ALL.exists():
        p_fig8 = insert_paragraph_before(
            anchor,
            "图 8 多算例对比汇总图（单轴/剪力墙/多片剪力墙/多层壳）\n"
            "Figure 8 Summary comparison across multiple cases",
            style="Normal",
        )
        insert_picture(p_fig8, FIG_ALL)

    insert_paragraph_before(anchor, "4.5 验证小结（表 3）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "表 3 汇总单元级材料测试、结构级对比校验（单轴+剪力墙）与软件级校核（sw1-1/sw2-1/Multi-layer Shell）的"
        "主要指标与结论。总体而言，新增对比算例与成图脚本强化了“可复现数据—可复核图形—可解释差异”的证据链条。",
        style="Normal",
    )

    tcap = insert_paragraph_before(
        anchor,
        "表 3 验证结果汇总\nTable 3 Summary of verification results",
        style="Normal",
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "类别"
    hdr[1].text = "算例/路径"
    hdr[2].text = "关键指标"
    hdr[3].text = "结论"

    def add_row(cat, case, metric, concl):
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = case
        row[2].text = metric
        row[3].text = concl

    add_row("单元级材料测试", "单轴压缩/拉伸/纯剪切/循环压拉/四阶段循环", "0 失败；软化/滞回路径符合预期", "通过")
    add_row("结构级对比校验", "单轴（压缩/拉伸/循环）", "峰值与软化段与参考吻合", "通过")
    add_row("结构级对比校验", "剪力墙循环加载", "滞回形态一致；峰值略低约 5%～10%", "通过")
    add_row("软件级校核", "sw1-1 多片剪力墙", "500/500 步，0 失败", "通过")
    add_row("软件级校核", "sw2-1 多片剪力墙", "≈452/500 步，后期奇异/收敛困难", "部分通过")
    add_row("软件级校核", "Multi-layer Shell", "947/947 步，0 失败", "通过")

    marker = anchor.insert_paragraph_before("")
    marker._element.addprevious(table._element)
    remove_paragraph(marker)

    insert_paragraph_before(anchor, "")


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    doc = Document(str(SRC))
    normalize_abstract(doc)
    rebuild_section_4(doc)
    doc.save(str(OUT))
    print("Wrote", OUT)


if __name__ == "__main__":
    main()

