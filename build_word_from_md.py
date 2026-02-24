import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MD_PATH = Path(r"e:\Basic\Concrete_Model\ADINA\论文正文_第1-3节_工程力学格式.md")
OUT_DOCX = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper.docx")
FIG1_PATH = Path(r"e:\Basic\Concrete_Model\ADINA\figure1_flow.png")
FIG3_PATH = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_uniaxial.png")
FIG4_PATH = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_wall.png")


def set_style_font(style, east_asia: str, ascii_font: str = "Times New Roman", size_pt: float = 10.5):
    font = style.font
    font.name = ascii_font
    font.size = Pt(size_pt)
    rfonts = style._element.rPr.rFonts
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def generate_flow_figure_png(path: Path):
    # Minimal flowchart as a PNG (matplotlib only).
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    fig = plt.figure(figsize=(10, 4), dpi=200)
    ax = fig.add_subplot(111)
    ax.set_axis_off()

    boxes = [
        (0.04, 0.55, 0.18, 0.28, "需求/规格\n（接口+参数+状态）"),
        (0.28, 0.55, 0.18, 0.28, "AI 生成/改写\nforumat.f90"),
        (0.52, 0.55, 0.18, 0.28, "standalone 测试驱动\n（78+ 项）"),
        (0.76, 0.55, 0.20, 0.28, "结构级校验\n（单轴+剪力墙）"),
        (0.40, 0.12, 0.30, 0.25, "失败→定位→修改→回归\n（Cursor 自动编译/执行/测试/修改）"),
    ]

    for x, y, w, h, txt in boxes:
        patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.02",
                               linewidth=1.2, edgecolor="#1f4e79", facecolor="#eaf2fb")
        ax.add_patch(patch)
        ax.text(x + w/2, y + h/2, txt, ha="center", va="center", fontsize=10)

    def arrow(p1, p2):
        ax.add_patch(FancyArrowPatch(p1, p2, arrowstyle="->", mutation_scale=12, linewidth=1.2, color="#1f4e79"))

    # main arrows
    arrow((0.22, 0.69), (0.28, 0.69))
    arrow((0.46, 0.69), (0.52, 0.69))
    arrow((0.70, 0.69), (0.76, 0.69))

    # feedback loop down and back to AI
    arrow((0.86, 0.55), (0.65, 0.37))
    arrow((0.40, 0.37), (0.35, 0.55))

    ax.text(0.5, 0.95, "图1 基于 AI 的本构集成开发方法框架 / Figure 1 Framework of AI-assisted constitutive integration and development",
            ha="center", va="center", fontsize=10, color="#000000")

    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def iter_blocks(md_text: str):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # horizontal rules
        if line.strip() in {"---", "—"}:
            i += 1
            continue

        # tables (markdown pipe tables)
        if line.strip().startswith("|") and "|" in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            yield ("table", table_lines)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            yield ("heading", (level, text))
            i += 1
            continue

        # blank lines
        if not line.strip():
            yield ("blank", "")
            i += 1
            continue

        # normal paragraph (merge consecutive non-blank non-table non-heading lines)
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6})\s+", nxt):
                break
            if nxt.strip().startswith("|") and "|" in nxt:
                break
            if nxt.strip() in {"---", "—"}:
                break
            para_lines.append(nxt.rstrip("\n"))
            i += 1
        yield ("para", "\n".join(para_lines).strip())


def add_markdown_table(doc: Document, table_lines):
    # Remove separator row like |---|---|
    rows = []
    for ln in table_lines:
        parts = [p.strip() for p in ln.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", p) for p in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(ncols):
            cell = table.cell(r, c)
            cell.text = row[c] if c < len(row) else ""


def main():
    if not MD_PATH.exists():
        raise FileNotFoundError(MD_PATH)

    # Generate Figure 1 if not present
    generate_flow_figure_png(FIG1_PATH)

    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()

    # Page setup (A4, typical margins)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    # Fonts
    set_style_font(doc.styles["Normal"], east_asia="宋体", size_pt=10.5)
    set_style_font(doc.styles["Title"], east_asia="黑体", size_pt=16)
    set_style_font(doc.styles["Subtitle"], east_asia="黑体", size_pt=12)
    for h in ["Heading 1", "Heading 2", "Heading 3"]:
        if h in doc.styles:
            set_style_font(doc.styles[h], east_asia="黑体", size_pt=12 if h == "Heading 1" else 11)

    # Simple MD → DOCX
    last_was_blank = True
    for kind, payload in iter_blocks(md_text):
        if kind == "blank":
            last_was_blank = True
            continue

        if kind == "heading":
            level, text = payload
            if level == 1:
                p = doc.add_paragraph(text, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                p = doc.add_paragraph(text, style="Subtitle")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 3:
                doc.add_heading(text, level=1)
            elif level == 4:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
            last_was_blank = False
            continue

        if kind == "table":
            add_markdown_table(doc, payload)
            doc.add_paragraph("")  # spacer
            last_was_blank = True
            continue

        if kind == "para":
            text = payload
            # strip bold markers (**...**) to plain text for Word
            text_plain = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            p = doc.add_paragraph(text_plain)
            last_was_blank = False

            # Insert figures right after captions
            if "图 1" in text_plain and "Figure 1" in text_plain and FIG1_PATH.exists():
                pic = doc.add_picture(str(FIG1_PATH), width=Cm(15.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if "图 3" in text_plain and "Figure 3" in text_plain and FIG3_PATH.exists():
                doc.add_picture(str(FIG3_PATH), width=Cm(15.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if "图 4" in text_plain and "Figure 4" in text_plain and FIG4_PATH.exists():
                doc.add_picture(str(FIG4_PATH), width=Cm(15.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(OUT_DOCX))
    print("Wrote", OUT_DOCX)


if __name__ == "__main__":
    main()

