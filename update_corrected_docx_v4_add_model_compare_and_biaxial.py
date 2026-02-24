from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.shared import Cm


SRC = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper_corrected_v3.docx")
OUT = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper_corrected_v4.docx")
BACKUP = Path(r"e:\Basic\Concrete_Model\ADINA\AI_VibeCoding_IndustrialFEM_Dev_Paper_corrected_v3.bak.docx")

FIG_UNIAX = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_uniaxial.png")
FIG_WALL = Path(r"e:\Basic\Concrete_Model\ADINA\verify\verify_wall.png")

FIG_SW1 = Path(r"e:\Basic\Concrete_Model\ADINA\compare_sw1-1.png")
FIG_SW2 = Path(r"e:\Basic\Concrete_Model\ADINA\compare_sw2-1.png")
FIG_ML = Path(r"e:\Basic\Concrete_Model\ADINA\compare_Multi-layer_Shell.png")
FIG_ALL = Path(r"e:\Basic\Concrete_Model\ADINA\all_cases_comparison_0818.png")

FIG_BIAX_PATHS = Path(r"e:\Basic\Concrete_Model\ADINA\standalone_tests\biaxial_paths_report.png")
FIG_BIAX_COMP = Path(r"e:\Basic\Concrete_Model\ADINA\standalone_tests\biaxial_comparison_report.png")


def remove_paragraph(paragraph: Paragraph) -> None:
    p: CT_P = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None  # type: ignore


def insert_paragraph_before(anchor: Paragraph, text: str, style: Optional[str] = None) -> Paragraph:
    p = anchor.insert_paragraph_before(text)
    if style:
        p.style = style
    return p


def insert_picture(anchor: Paragraph, image_path: Path, width_cm: float = 15.5) -> None:
    # Insert picture ABOVE caption (caption paragraph is anchor)
    pic_p = anchor.insert_paragraph_before("")
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_first_paragraph_index(doc: Document, startswith: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    raise ValueError(f"Cannot find paragraph starting with: {startswith!r}")


def normalize_abstract_add_biaxial(doc: Document) -> None:
    """
    Append one sentence about biaxial verification (if not already present).
    """
    add = "并补充双向应力路径与双轴强度对照测试，分析平面应力条件下双轴增强来源与实验偏差。"
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t.startswith("摘要："):
            continue
        if "双向应力" in t or "双轴强度" in t:
            return
        p.text = t + add
        return


def update_data_availability(doc: Document) -> None:
    """
    Extend the '数据与代码可获得性说明' paragraph to include new biaxial scripts.
    """
    add = "此外，新增双向应力路径测试脚本位于 standalone_tests/test_biaxial_paths.f90，成图脚本为 " \
          "standalone_tests/plot_biaxial_paths.py 与 standalone_tests/plot_biaxial_comparison.py。"
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("为保证研究的可重复性"):
            if "。；新增双向应力路径测试脚本" in t:
                p.text = t.replace("。；新增双向应力路径测试脚本", "。此外，新增双向应力路径测试脚本")
                return
            if "test_biaxial_paths" in t:
                # already added (but ensure punctuation)
                if t.endswith("comparison.py"):
                    p.text = t + "。"
                return
            # add as a new sentence
            if not t.endswith("。"):
                t = t + "。"
            p.text = t + add
            return


def insert_section_3_5_model_compare(doc: Document) -> None:
    """
    Insert a new subsection 3.5 before Section 4.
    """
    i4 = find_first_paragraph_index(doc, "4 ")
    anchor = doc.paragraphs[i4]

    # Guard: avoid duplicate insertion
    for p in doc.paragraphs[max(i4 - 30, 0):i4]:
        if p.text.strip().startswith("3.5 ") and ("论文模型" in p.text or "Bathe" in p.text):
            return

    insert_paragraph_before(
        anchor,
        "3.5 论文模型的机理理解、平面应力实现特征与差异对比",
        style="Heading 1",
    )

    insert_paragraph_before(
        anchor,
        "Bathe-Ramaswamy（1979）模型以单轴 Saenz 曲线为基础，通过“加载函数—等效模量—破坏包络”三者耦合来刻画多轴应力下的非线性。"
        "在加载阶段，模型先计算主应力与主应力方向，并分别沿三条主方向建立等效单轴切线模量，再通过权重规则形成等效各向同性模量或三向正交异性模量［1］。"
        "同时采用 Kupfer 等实验为基础的压碎包络面与抗拉包络面（图 2～图 3）来识别开裂/压碎，并用式（20）的强度放大系数 \\(\\gamma_1\\) 将多轴应力效应反馈到单轴强度参数（\\(\\sigma_c,\\,\\sigma_u,\\,\\varepsilon_c,\\,\\varepsilon_u\\)）中，从而得到多轴条件下的“等效 Saenz 曲线”［1，4，9］。",
        style="Normal",
    )

    insert_paragraph_before(
        anchor,
        "在 OpenSees 平面应力用户材料（PSUMAT）接口下，本文实现以 ADINA（MODEL=5 的 2D 版本）为蓝本复现其分支结构（压碎/裂缝/未开裂）。"
        "与论文的三维推导相比，平面应力约束 \\(\\sigma_3=0\\) 使压碎包络查表（SP1、SP31～SP33）在许多应力状态下退化：排序后的最大主应力常为 \\(P_1=0\\)，"
        "从而对应的查表位置为 \\(\\sigma_1/\\sigma_c=0\\)，插值结果给出 \\(SP_{31}=SP_{32}=SP_{33}=1\\)，即 \\(\\gamma_1=1\\)。"
        "因此，在未开裂压缩阶段（Branch C）若按总应变割线法更新（\\(\\sigma=C(E_{sec})\\,\\varepsilon\\)），双轴强度增强主要来自平面应力本构矩阵中的 Poisson 耦合项，而非包络面反馈；这一点与论文模型在 2D 平面应力退化情形下的行为是一致的。",
        style="Normal",
    )

    insert_paragraph_before(
        anchor,
        "另一方面，本文在工程可用性与数值稳定性约束下对原模型做了若干“实现层”的补充与修正："
        "（1）为避免裂缝开合过程的滞弹性漂移，在闭合阶段引入渐进混合法（hypoelastic-total blend）；"
        "（2）为避免循环加载下方向切线模量污染 Saenz 参数，采用关键模块变量的保存/恢复机制，消除循环压缩应力异常超越单调峰值的问题；"
        "（3）在拉伸侧引入线性软化与抗拉强度下限保护以增强鲁棒性。"
        "这些处理不改变论文模型的核心物理假设（固定裂缝、压碎/开裂包络、Saenz 压缩曲线），但显著改善了 NR 收敛与工程算例可复现性，相关验证见第 4 节。",
        style="Normal",
    )


def rebuild_section_4_with_biaxial(doc: Document) -> None:
    """
    Rebuild Section 4 and insert new biaxial verification subsection + figures.
    Keep existing figure numbers 3-8, add Figure 9-10.
    """
    i4 = find_first_paragraph_index(doc, "4 ")
    i5 = find_first_paragraph_index(doc, "5 ")

    # remove old section 4 content
    for p in list(doc.paragraphs[i4:i5])[::-1]:
        remove_paragraph(p)

    anchor = doc.paragraphs[find_first_paragraph_index(doc, "5 ")]

    insert_paragraph_before(anchor, "4 验证设计及结果", style="Subtitle")

    insert_paragraph_before(anchor, "4.1 验证体系与评价指标", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "验证采用“单元级—结构级—软件级”三级体系。单元级通过 standalone Fortran 程序直接调用 PSUMAT 接口，"
        "在不依赖 OpenSees 可执行文件的前提下开展回归测试；结构级将 forumat.f90 编译进 OpenSees 后，"
        "采用与参考模型一致的算例与加载协议对比响应；软件级进一步选取多片剪力墙与多层壳单元等算例，"
        "检验大变形循环加载下的收敛性与数值稳定性。剪力墙算例中，基底剪力取底部约束节点反力之和 "
        "V=ΣRi(i=1…5)，控制位移取加载点水平位移 u，比较 V-u 滞回曲线峰值、对称性与收敛步数等指标。",
        style="Normal",
    )

    insert_paragraph_before(anchor, "4.2 单元级材料测试（5个）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "依据最新开发记录，单元级材料测试覆盖5条典型路径：单轴压缩（至 ε=-0.003）、单轴拉伸（至 ε=+0.0005）、"
        "纯剪切（γ=0.005）、循环压拉（压缩→卸载→拉伸→再压缩）与四阶段循环（压→卸→拉→再压）。"
        "上述测试均顺利完成且无失败步，为后续结构级与软件级算例提供了可追溯的单元级正确性保证。",
        style="Normal",
    )

    insert_paragraph_before(anchor, "4.3 结构级对比校验（单轴+剪力墙）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "结构级校验聚焦两类最具代表性的平面应力应用场景：单轴（压缩/拉伸/循环）与悬臂剪力墙循环加载。"
        "对比结果表明：B&R 与参考模型在单轴峰值与软化段吻合良好；剪力墙滞回环形态一致，B&R 峰值底部剪力"
        "略低约 5%～10%，包络线与退化趋势吻合。对应对比图见图3～图4。",
        style="Normal",
    )

    p_fig3 = insert_paragraph_before(
        anchor,
        "图 3 单轴压缩与拉伸应力-应变曲线（集成模型与参考模型）\n"
        "Figure 3 Uniaxial compressive and tensile stress-strain curves (integrated model vs reference)",
        style="Normal",
    )
    if FIG_UNIAX.exists():
        insert_picture(p_fig3, FIG_UNIAX)

    p_fig4 = insert_paragraph_before(
        anchor,
        "图 4 剪力墙侧向力-位移滞回曲线（集成模型与参考模型）\n"
        "Figure 4 Shear wall lateral force-displacement hysteresis (integrated model vs reference)",
        style="Normal",
    )
    if FIG_WALL.exists():
        insert_picture(p_fig4, FIG_WALL)

    insert_paragraph_before(anchor, "4.4 多片剪力墙与多层壳单元软件级校核（含对比图与成图脚本）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "为检验本构在更贴近工程使用情景（大变形+循环+耦合单元/求解器）的稳定性，进一步选取多片剪力墙"
        "（sw1-1、sw2-1）与多层壳单元（Multi-layer Shell）算例开展软件级校核。最新记录显示：sw1-1 可完成"
        "500/500 步、0 失败；sw2-1 在大位移后期完成约 90% 进度后出现矩阵奇异相关的收敛困难；Multi-layer Shell"
        "可完成 947/947 步、0 失败。三算例与参考模型的对比图分别见图5～图7。",
        style="Normal",
    )
    p_fig5 = insert_paragraph_before(
        anchor,
        "图 5 sw1-1 剪力墙对比（集成模型与参考模型）\n"
        "Figure 5 sw1-1 shear wall comparison (integrated model vs reference)",
        style="Normal",
    )
    if FIG_SW1.exists():
        insert_picture(p_fig5, FIG_SW1)

    p_fig6 = insert_paragraph_before(
        anchor,
        "图 6 sw2-1 剪力墙对比（集成模型与参考模型）\n"
        "Figure 6 sw2-1 shear wall comparison (integrated model vs reference)",
        style="Normal",
    )
    if FIG_SW2.exists():
        insert_picture(p_fig6, FIG_SW2)

    p_fig7 = insert_paragraph_before(
        anchor,
        "图 7 Multi-layer Shell 对比（集成模型与参考模型）\n"
        "Figure 7 Multi-layer Shell comparison (integrated model vs reference)",
        style="Normal",
    )
    if FIG_ML.exists():
        insert_picture(p_fig7, FIG_ML)

    p_fig8 = insert_paragraph_before(
        anchor,
        "图 8 多算例对比汇总图（单轴/剪力墙/多片剪力墙/多层壳）\n"
        "Figure 8 Summary comparison across multiple cases",
        style="Normal",
    )
    if FIG_ALL.exists():
        insert_picture(p_fig8, FIG_ALL)

    # ------------------------------------------------------------------
    # New biaxial subsection + figures
    # ------------------------------------------------------------------
    insert_paragraph_before(anchor, "4.5 双向应力路径与双轴强度校核（对照论文机理）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "为补充平面应力本构在双向应力状态下的行为验证，构建 8 条比例应力路径：-1:-1、-1:-0.5、-1:-0.2、"
        "-1:0、-1:0.1、-1:0.2、-1:0.5 与 1:1。测试采用 standalone 程序在比例应变加载下调用 PSUMAT，记录"
        "各路径的 \\(\\sigma_1/f_c\\)、\\(\\sigma_2/f_c\\) 峰值及完整应力-应变曲线。结果表明：等双轴受压路径的峰值"
        "\\(\\sigma_1/f_c\\approx\\sigma_2/f_c\\approx 1/(1-\\nu)=1.25\\)；当 \\(\\sigma_2/\\sigma_1\\) 从 1 减小到 0 时，"
        "增强效应逐步消退并回归单轴 \\(\\sigma_1/f_c=1\\)。这一趋势符合论文模型在平面应力退化时“包络面反馈弱、"
        "双轴增强主要由 Poisson 耦合贡献”的机理判断（见第 3.5 节）。与 Kupfer 等实验结果［4］相比，等双轴受压"
        "峰值偏高约 7%～8%，反映了常 Poisson 比假设与 2D 退化包络对真实双轴强度提升的有限表达能力。"
        "双向路径响应与强度包络对比见图9～图10。",
        style="Normal",
    )

    p_fig9 = insert_paragraph_before(
        anchor,
        "图 9 双向应力路径下的应力-应变响应与应力轨迹（8 条路径）\n"
        "Figure 9 Stress-strain responses and stress trajectories under biaxial stress paths (8 paths)",
        style="Normal",
    )
    if FIG_BIAX_PATHS.exists():
        insert_picture(p_fig9, FIG_BIAX_PATHS)

    p_fig10 = insert_paragraph_before(
        anchor,
        "图 10 双轴强度包络对比：本文实现、论文平面应力退化机理与 Kupfer 实验数据［4］\n"
        "Figure 10 Biaxial strength envelope comparison: this implementation, plane-stress degenerated mechanism and Kupfer data [4]",
        style="Normal",
    )
    if FIG_BIAX_COMP.exists():
        insert_picture(p_fig10, FIG_BIAX_COMP)

    # ------------------------------------------------------------------
    # Summary + table (renumbered to 4.6)
    # ------------------------------------------------------------------
    insert_paragraph_before(anchor, "4.6 验证小结（表 3）", style="Heading 1")
    insert_paragraph_before(
        anchor,
        "表 3 汇总单元级材料测试、结构级对比校验（单轴+剪力墙）、软件级校核（sw1-1/sw2-1/Multi-layer Shell）"
        "以及新增的双向应力路径与双轴强度校核的主要指标与结论。总体而言，新增双轴测试补齐了“单轴—双轴—"
        "结构响应”的证据链条，并增强了对论文模型机理与平面应力退化特征的可解释性。",
        style="Normal",
    )

    insert_paragraph_before(
        anchor,
        "表 3 验证结果汇总\nTable 3 Summary of verification results",
        style="Normal",
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "类别"
    hdr[1].text = "算例/路径"
    hdr[2].text = "关键指标"
    hdr[3].text = "结论"

    def add_row(cat: str, case: str, metric: str, concl: str) -> None:
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = case
        row[2].text = metric
        row[3].text = concl

    add_row("单元级材料测试", "单轴压缩/拉伸/纯剪切/循环压拉/四阶段循环", "0 失败；软化/滞回路径符合预期", "通过")
    add_row("单元级材料测试", "双向应力路径（8条）", "峰值/包络趋势与平面应力机理一致；无数值发散", "通过")
    add_row("结构级对比校验", "单轴（压缩/拉伸/循环）", "峰值与软化段与参考吻合", "通过")
    add_row("结构级对比校验", "剪力墙循环加载", "滞回形态一致；峰值略低约 5%～10%", "通过")
    add_row("软件级校核", "sw1-1 多片剪力墙", "500/500 步，0 失败", "通过")
    add_row("软件级校核", "sw2-1 多片剪力墙", "≈452/500 步，后期奇异/收敛困难", "部分通过")
    add_row("软件级校核", "Multi-layer Shell", "947/947 步，0 失败", "通过")

    # Place table right after caption: use marker trick
    marker = anchor.insert_paragraph_before("")
    marker._element.addprevious(table._element)
    remove_paragraph(marker)

    insert_paragraph_before(anchor, "")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if not BACKUP.exists():
        BACKUP.write_bytes(SRC.read_bytes())

    doc = Document(str(SRC))

    normalize_abstract_add_biaxial(doc)
    insert_section_3_5_model_compare(doc)
    rebuild_section_4_with_biaxial(doc)
    update_data_availability(doc)

    doc.save(str(OUT))
    print("Wrote", OUT)
    print("Backup", BACKUP)


if __name__ == "__main__":
    main()

